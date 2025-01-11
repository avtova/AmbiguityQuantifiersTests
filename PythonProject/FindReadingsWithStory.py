
import pandas as pd

from openai import OpenAI
from docx import Document

# Set your OpenAI API key
API_KEY = ""
CHAT_GPT_VERSION = "gpt-3.5-turbo-16k" # "gpt-4o"
INPUT_DATA_EXCEL_FILE = "SeminarData.xlsx"
OUTPUT_FILE = "output_readings_with_story " + CHAT_GPT_VERSION + ".docx"

def main():
    # Load the Excel file
    data = load_and_clean_data (INPUT_DATA_EXCEL_FILE)

    # Create a Word document to store the results
    output_doc = Document()

    for i in range(4):
        # Process each row in the Excel file
        for index, row in data.iterrows():
            print(row['story_file_name'], row['sentence'])  # Replace with normalized names
            story_file_name = row['story_file_name']
            sentence = row['sentence']

            # Read and construct the prompt from the Word file
            try:
                story_doc = Document(story_file_name)
                story_text = "\n".join(paragraph.text for paragraph in story_doc.paragraphs)
            except Exception as e:
                print(f"Error reading Word file {story_file_name}: {e}")
                continue

            # generate readings when the story is not given
            result_text = handle_generate_readings_story_given(sentence, story_text, API_KEY)
            #print(result_text)

            # Write the result into the output Word file
            output_doc.add_heading(sentence, level=1)
            output_doc.add_paragraph(result_text)

            # Save the output Word document
        output_doc.save(OUTPUT_FILE)

def load_and_clean_data(file_path):
    """Load Excel data and clean column names."""
    data = pd.read_excel(file_path)
    data.columns = data.columns.str.lower()
    print("Cleaned column names:", data.columns.tolist())
    return data

def handle_generate_readings_story_given(story_text, sentence, api_key):
    intro = "Below is a story followed by an ambiguous sentence. Identify both the surface and inverse readings of the sentence, choose the reading that aligns best with the story and explain your reasoning."
    prompt = f"{intro}\n\n{story_text}\n\n{sentence}"
    return call_openai_api(prompt, api_key)

# Function definition for call_openai_api
def call_openai_api(prompt, api_key):
    """Function to make the API call."""
    try:
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model = CHAT_GPT_VERSION,
            messages=[
                {"role": "system",
                 "content": "You are a helpful assistant who answers common-sense reasoning questions."},
                {"role": "user", "content": prompt},
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error calling OpenAI API: {e}")
        return f"Error: {e}"

if __name__ == "__main__":
    main()